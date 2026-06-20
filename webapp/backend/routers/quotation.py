import math
import os
import re
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel

APP_DIR = Path(__file__).parent.parent.parent.parent
BACKEND_DIR = Path(__file__).parent.parent
TEMP_DB = str(APP_DIR / "temp.db")
TEMPLATES = APP_DIR / "templates"

sys.path.insert(0, str(APP_DIR))
sys.path.insert(0, str(BACKEND_DIR))
from tempquotebase import (
    init_temp_db, add_new_quote, get_all_quotes, delete_quote,
    add_product_quote, get_all_quote_products, clear_quotedata_table,
    get_highest_sr_no, get_db_connection, get_items_table_name,
)
from auth import get_current_user
from user_db import get_user_costing_db, get_user_temp_db, get_user_wizard_temp_db, get_user_sizing_db, get_user_inquiry_db
from sql_handler import fetch_sizing_by_sr

init_temp_db()
router = APIRouter()

def _tdb(username: str, scope: str) -> str:
    db = get_user_wizard_temp_db(username) if scope == "wizard" else get_user_temp_db(username)
    init_temp_db(db)
    return db

TEMPLATE_MAP = {
    "High voltage":                    "Quote_format_High_Vtg.docx",
    "Low voltage":                     "Quote_format_Low_Vtg.docx",
    "Extended Warranty High Voltage":  "Quote_format_Extended_Warranty_High_Vtg.docx",
    "Extended Warranty Low Voltage":   "Quote_format_Extended_Warranty_Low_Vtg.docx",
    "Low & High Voltage Export":       "Quote_format_Low_High_Vtg_Export.docx",
}
FORMATS = list(TEMPLATE_MAP.keys())

MODULAR_RACKS = {
    "W=600*D=1000*H=880":  30000,
    "W=600*D=1000*H=1392": 40000,
    "W=600*D=1000*H=1882": 49000,
    "W=600*D=1000*H=1971": 64000,
    "W=600*D=1000*H=2058": 69000,
    "W=600*D=800*H=992":   30000,
    "W=600*D=800*H=1704":  43000,
    "W=600*D=1000*H=2325": 70000,
    "W=600*D=1400*H=1882": 70000,
}

# ── Pydantic ──────────────────────────────────────────────────────────────────

class QuoteCreate(BaseModel):
    code: str
    date: str
    customer_name: str
    solution_provider: str
    format_name: str
    sales_person: str = ""

class QuoteItem(BaseModel):
    code: str
    format: str
    date: str
    solution_provider: str
    customer_name: str
    sr_no: int
    sol_no: int
    ups_rating: str = "-"
    backup_requirement: str = "-"
    calc_load: str = "-"
    celltype: str = "-"
    centre_tapping: str = "-"
    batterypartcode: str = "-"
    backup_time: str = "0"
    quantity: int = 1
    quote_price: float = 0
    modular_rack: str = "-"

class AddFromCostingReq(BaseModel):
    quote_code: str
    costing_row_index: int
    price_option: str   # "A","B-5","B","B+5","C","C+5","custom"
    quantity: int
    custom_pct: float = 0.0  # used when price_option == "custom"
    sizing_project: str = ""
    sizing_sr_no: int = 0

class AddModularReq(BaseModel):
    quote_code: str
    rack_key: str
    quantity: int
    custom_price: float = 0.0

# ── helpers ───────────────────────────────────────────────────────────────────

def _row_to_dict(row: tuple) -> dict:
    keys = ["code","format","date","solution_provider","customer_name",
            "sr_no","sol_no","ups_rating","backup_requirement","calc_load",
            "celltype","centre_tapping","batterypartcode","backup_time",
            "quantity","quote_price","modular_rack","system_text","solution_text",
            "calc_load_unit","item_type","ageing_type"]
    d = dict(zip(keys, row))
    d.setdefault("system_text", None)
    d.setdefault("solution_text", None)
    d.setdefault("calc_load_unit", "kW")
    d.setdefault("ageing_type", "BOL")
    return d


def _inr(amount: float) -> str:
    """Format a number in Indian numbering (e.g. 1,23,456.00)."""
    rounded = round(amount, 2)
    s = f"{rounded:.2f}"
    integer_part, decimal_part = s.split(".")
    neg = integer_part.startswith("-")
    if neg:
        integer_part = integer_part[1:]
    if len(integer_part) <= 3:
        formatted = integer_part
    else:
        last3 = integer_part[-3:]
        rest = integer_part[:-3]
        groups = []
        while rest:
            groups.append(rest[-2:] if len(rest) >= 2 else rest)
            rest = rest[:-2] if len(rest) > 2 else ""
        groups.reverse()
        formatted = ",".join(groups) + "," + last3
    return ("-" if neg else "") + formatted + "." + decimal_part


def _generate_docx(quote_code: str, db_path: str = None) -> str:
    import docx
    from docx.shared import Pt
    from docx.oxml.ns import qn

    quotes = get_all_quotes(db_path)
    meta = next((q for q in quotes if q[0] == quote_code), None)
    if not meta:
        raise ValueError("Quote not found")
    code, date, customer, provider, fmt_fname, *_ = meta
    template_path = str(TEMPLATES / fmt_fname)
    items = get_all_quote_products(quote_code, db_path)

    doc = docx.Document(template_path)
    replacements = {
        "{{CODE}}": str(code),
        "{{DATE}}": str(date),
        "{{SOLUTION_PROVIDER}}": str(provider),
        "{{CUSTOMER_NAME}}": str(customer),
    }
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, val)
        for run in para.runs:
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
            run.font.size = Pt(11)

    if not doc.tables:
        raise ValueError("No table in template")
    table = doc.tables[0]

    sol_no = 0
    for row_idx, item in enumerate(items, start=1):
        d = _row_to_dict(item)
        while len(table.rows) <= row_idx:
            table.add_row()
        trow = table.rows[row_idx]

        if d["modular_rack"] and d["modular_rack"] != "-":
            system_text = ""
            rack_key = d["modular_rack"]
            solution_text = f"Modular Battery Rack ({rack_key})" if rack_key.startswith("W=") else rack_key
        else:
            sol_no += 1
            bt = d["backup_time"]
            try:
                bt_floor = str(math.floor(float(bt))) if bt and bt not in ("-", "0", "") else "-"
            except Exception:
                bt_floor = "-"
            ageing_lbl = d.get("ageing_type") or "BOL"
            load_unit = d.get("calc_load_unit") or "kW"
            load_line = f"\n(Load: {d['calc_load']}{load_unit})" if d.get("calc_load") else ""
            system_text = d.get("system_text") or (
                f"{d['ups_rating']}KVA : {d['backup_requirement']}Min Backup"
                f"{load_line}\n"
                f"(Cell Type:{d['celltype']})\n"
                f"({d['centre_tapping']})"
            )
            solution_text = d.get("solution_text") or (
                f"Solution{sol_no}: Lithium Battery Pack\n"
                f"({d['batterypartcode']}) with\n"
                f"Approximate Backup Time: {bt_floor}Mins At {ageing_lbl}\n"
                f"With Cabinet and inbuilt BMS"
            )

        cells = trow.cells
        def _set(cell, text):
            cell.text = str(text)
        if len(cells) > 0: _set(cells[0], row_idx)
        if len(cells) > 1: _set(cells[1], system_text)
        if len(cells) > 2: _set(cells[2], solution_text)
        if len(cells) > 3: _set(cells[3], d["quantity"])
        try:
            price_f = float(str(d["quote_price"]).replace("Rs.", "").replace("/- +GST", "").replace(",", "").strip())
        except Exception:
            price_f = 0
        if len(cells) > 4:
            _set(cells[4], f"Rs. {_inr(price_f)}/- +GST")
        total = (int(d["quantity"]) if str(d["quantity"]).isdigit() else 0) * price_f
        if len(cells) > 5:
            _set(cells[5], f"Rs. {_inr(total)}/- +GST")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


# ── quote CRUD ────────────────────────────────────────────────────────────────

@router.get("/formats")
def list_formats(_=Depends(get_current_user)):
    return FORMATS


@router.get("/modular-racks")
def list_racks(_=Depends(get_current_user)):
    return [{"key": k, "price": v} for k, v in MODULAR_RACKS.items()]


@router.get("/next-code")
def next_code(user=Depends(get_current_user), scope: str = Query("regular")):
    tdb = _tdb(user["username"], scope)
    try:
        from basefunctions import get_new_quote_code
        return {"code": get_new_quote_code()}
    except Exception:
        rows = get_all_quotes(tdb)
        numeric = [int(r[0]) for r in rows if str(r[0]).isdigit()]
        return {"code": str(max(numeric) + 1) if numeric else "1"}


@router.get("/quotes")
def list_quotes(user=Depends(get_current_user), scope: str = Query("regular")):
    tdb = _tdb(user["username"], scope)
    rows = get_all_quotes(tdb)
    return [{"code": r[0], "date": r[1], "customer_name": r[2],
             "solution_provider": r[3], "format": r[4],
             "sales_person": r[5] if len(r) > 5 else ""} for r in rows]


@router.post("/quotes", status_code=201)
def create_quote(body: QuoteCreate, user=Depends(get_current_user), scope: str = Query("regular")):
    tdb = _tdb(user["username"], scope)
    fname = TEMPLATE_MAP.get(body.format_name, "Quote_format_High_Vtg.docx")
    try:
        add_new_quote(body.code, body.date, body.customer_name, body.solution_provider, fname, tdb, body.sales_person)
    except Exception as e:
        raise HTTPException(500, str(e))
    return {"code": body.code}


class PatchMetaReq(BaseModel):
    customer_name: str = ""
    solution_provider: str = ""
    sales_person: str = ""
    date: str = ""
    format_name: str = ""
    new_code: str = ""

@router.patch("/quotes/{code}/meta", status_code=200)
def patch_meta(code: str, body: PatchMetaReq, user=Depends(get_current_user), scope: str = Query("regular")):
    import sqlite3 as _sq3
    tdb = _tdb(user["username"], scope)
    quotes = get_all_quotes(tdb)
    meta = next((q for q in quotes if q[0] == code), None)
    if not meta:
        raise HTTPException(404, "Quote not found")

    fname = TEMPLATE_MAP.get(body.format_name, meta[4]) if body.format_name else meta[4]
    new_code = body.new_code.strip() if body.new_code.strip() else code
    customer = body.customer_name or meta[2]
    provider = body.solution_provider or meta[3]
    sales = body.sales_person if body.sales_person is not None else (meta[5] if len(meta) > 5 else "")
    date = body.date or meta[1]

    conn = get_db_connection(tdb)
    try:
        c = conn.cursor()
        c.execute(
            "UPDATE active_quotes SET customer_name=?, solution_provider=?, sales_person=?, date=?, format=? WHERE code=?",
            (customer, provider, sales, date, fname, code)
        )
        if new_code != code:
            existing = c.execute("SELECT code FROM active_quotes WHERE code=?", (new_code,)).fetchone()
            if existing:
                raise HTTPException(400, f"Quote code '{new_code}' already exists")
            c.execute("UPDATE active_quotes SET code=? WHERE code=?", (new_code, code))
            old_tbl = get_items_table_name(code)
            new_tbl = get_items_table_name(new_code)
            try:
                c.execute(f'ALTER TABLE "{old_tbl}" RENAME TO "{new_tbl}"')
                c.execute(f'UPDATE "{new_tbl}" SET code=?', (new_code,))
            except Exception:
                pass
        conn.commit()
    finally:
        conn.close()

    try:
        from inquiry_db import _conn as _inq_conn, init_inquiry_db as _inq_init
        user_inq_db = get_user_inquiry_db(user["username"])
        for _dbp in [user_inq_db, None]:
            _inq_init(_dbp)
            with _inq_conn(_dbp) as c:
                c.execute(
                    "UPDATE inquiry SET project_customer=?, solution_provider=?, sales_person=? WHERE quote_code=?",
                    (customer, provider, sales, code)
                )
                if new_code != code:
                    c.execute("UPDATE inquiry SET quote_code=? WHERE quote_code=?", (new_code, code))
    except Exception:
        pass

    return {"detail": "updated", "new_code": new_code}


@router.delete("/quotes/{code}")
def remove_quote(code: str, user=Depends(get_current_user), scope: str = Query("regular")):
    tdb = _tdb(user["username"], scope)
    try:
        delete_quote(code, tdb)
    except Exception as e:
        raise HTTPException(500, str(e))
    try:
        from inquiry_db import _conn as _inq_conn, init_inquiry_db as _inq_init
        user_inq_db = get_user_inquiry_db(user["username"])
        _inq_init(user_inq_db)
        with _inq_conn(user_inq_db) as c:
            c.execute('DELETE FROM inquiry WHERE quote_code = ?', (code,))
    except Exception:
        pass
    return {"detail": "deleted"}


@router.get("/quotes/{code}/items")
def get_items(code: str, user=Depends(get_current_user)):
    tdb = get_user_temp_db(user["username"])
    rows = get_all_quote_products(code, tdb)
    return [_row_to_dict(r) for r in rows]


@router.post("/quotes/{code}/items", status_code=201)
def add_item(code: str, body: QuoteItem, user=Depends(get_current_user)):
    tdb = get_user_temp_db(user["username"])
    try:
        add_product_quote(
            code, body.code, body.format, body.date, body.solution_provider,
            body.customer_name, body.sr_no, body.sol_no, body.ups_rating,
            body.backup_requirement, body.calc_load, body.celltype,
            body.centre_tapping, body.batterypartcode, body.backup_time,
            body.quantity, body.quote_price, body.modular_rack, db_path=tdb,
        )
    except Exception as e:
        raise HTTPException(500, str(e))
    return {"detail": "added"}


@router.delete("/quotes/{code}/items/{sr_no}")
def delete_item(code: str, sr_no: int, user=Depends(get_current_user)):
    tdb = get_user_temp_db(user["username"])
    all_products = get_all_quote_products(code, tdb)
    quotes = get_all_quotes(tdb)
    meta = next((q for q in quotes if q[0] == code), None)
    if not meta:
        raise HTTPException(404, "Quote not found")
    _, date, customer, provider, fmt, *_ = meta

    deleted_d = next((_row_to_dict(i) for i in all_products if str(_row_to_dict(i)["sr_no"]) == str(sr_no)), None)

    clear_quotedata_table(code, tdb)
    new_sr = 0
    for item in all_products:
        d = _row_to_dict(item)
        if str(d["sr_no"]) == str(sr_no):
            continue
        new_sr += 1
        add_product_quote(
            code, d["code"], d["format"], d["date"], d["solution_provider"],
            d["customer_name"], new_sr, d["sol_no"], d["ups_rating"],
            d["backup_requirement"], d["calc_load"], d["celltype"],
            d["centre_tapping"], d["batterypartcode"], d["backup_time"],
            d["quantity"], d["quote_price"], d["modular_rack"],
            item_type=d.get("item_type") or "system", ageing_type=d.get("ageing_type") or "BOL", db_path=tdb,
        )
    try:
        from inquiry_db import sync_inquiry_for_quote as _sync_inq, _conn as _inq_conn, init_inquiry_db as _inq_init
        user_inq_db = get_user_inquiry_db(user["username"])
        _inq_init(user_inq_db)
        if deleted_d and str(deleted_d.get("item_type", "system")) == "system":
            with _inq_conn(user_inq_db) as c:
                c.execute('DELETE FROM inquiry WHERE quote_code = ? AND sol_no = ?',
                          (code, str(deleted_d.get("sol_no", ""))))
        updated = [_row_to_dict(i) for i in get_all_quote_products(code, tdb)]
        _sync_inq(code, updated, db_path=user_inq_db)
    except Exception:
        pass
    return {"detail": "deleted"}


# ── add from costing ──────────────────────────────────────────────────────────

@router.post("/quotes/{code}/add-from-costing", status_code=201)
def add_from_costing(code: str, body: AddFromCostingReq, user=Depends(get_current_user)):
    import sqlite3
    costing_db = get_user_costing_db(user["username"])
    try:
        conn = sqlite3.connect(costing_db)
        cur = conn.cursor()
        cur.execute("PRAGMA table_info(tree)")
        cols = [r[1] for r in cur.fetchall()]
        cur.execute("SELECT * FROM tree")
        rows = cur.fetchall()
        conn.close()
    except Exception as e:
        raise HTTPException(500, f"Costing DB error: {e}")

    if body.costing_row_index >= len(rows):
        raise HTTPException(404, "Costing row not found")

    row = rows[body.costing_row_index]

    def idx(name):
        try: return cols.index(name)
        except ValueError: return -1

    base_price = float(row[idx("Total Cost of Pack (A)")] or 0)
    multipliers = {"A": 1.0, "B-5": 1.05, "B": 1.10, "B+5": 1.15, "C": 1.20, "C+5": 1.25}
    if body.price_option == "custom":
        mult = 1.0 + (body.custom_pct / 100.0)
    else:
        mult = multipliers.get(body.price_option, 1.0)
    quote_price = round(base_price * mult, 2)

    duration_raw = str(row[idx("Duration")] or "0")
    try:
        backup_time = float("".join(c for c in duration_raw if c.isdigit() or c == "."))
    except Exception:
        backup_time = 0.0

    celltype = row[idx("Cylindrical/ Prismatic")] or "-"
    centre_tapping = row[idx("Centre tap/non centre tap")] or "-"
    batterypartcode = row[idx("Battery Partcode")] or "-"

    tdb = get_user_temp_db(user["username"])
    quotes = get_all_quotes(tdb)
    meta = next((q for q in quotes if q[0] == code), None)
    if not meta:
        raise HTTPException(404, "Quote not found")
    q_code, q_date, q_customer, q_provider, q_fmt, *_rest = meta; q_sales = _rest[0] if _rest else ""

    sr_no = get_highest_sr_no(code, tdb) + 1
    all_items = get_all_quote_products(code, tdb)
    sol_no = sum(1 for i in all_items if str(_row_to_dict(i)["modular_rack"]) == "-") + 1

    add_product_quote(
        code, q_code, q_fmt, q_date, q_provider, q_customer,
        sr_no, sol_no, "-", "-", "-",
        str(celltype), str(centre_tapping), str(batterypartcode),
        "-", body.quantity, quote_price, "-", item_type="system", ageing_type="BOL", db_path=tdb,
    )

    if body.sizing_project and body.sizing_sr_no:
        try:
            import time as _time
            from inquiry_db import push_row as _push_inq
            _yr = _time.localtime().tm_year % 100
            _type = f"EVTPL/{_yr:02d}-{(_yr+1):02d}/{code}"
            sdb = get_user_sizing_db(user["username"])
            srow = fetch_sizing_by_sr(body.sizing_project, body.sizing_sr_no, db_path=sdb)
            user_inq_db = get_user_inquiry_db(user["username"])
            if srow:
                unit_price = round(quote_price / body.quantity, 2) if body.quantity else quote_price
                _push_inq({
                    "inquiry_date": q_date,
                    "type": _type, "sales_person": str(q_sales or ""),
                    "solution_provider": str(q_provider or ""),
                    "project_customer": str(q_customer or ""),
                    "ups_make": str(srow[3] or ""), "ups_model": str(srow[4] or ""),
                    "ups_kva": str(srow[5] or ""),
                    "actual_load_kva": str(srow[6] or ""), "load_kw": str(srow[7] or ""),
                    "power_factor": str(srow[8] or ""), "inverter_efficiency": str(srow[9] or ""),
                    "dc_voltage": str(srow[10] or ""), "backup_min": str(srow[11] or ""),
                    "cell_chemistry": str(srow[17] or ""),
                    "ageing_pct": str(srow[12] or ""), "design_margin_pct": str(srow[13] or ""),
                    "dod_margin_pct": str(srow[14] or ""), "derating_pct": str(srow[15] or ""),
                    "capacity_ah": str(srow[27] or ""),
                    "centre_tap": str(centre_tapping or ""), "cell_type": str(celltype or ""),
                    "ageing_type": "BOL", "backup_time_min": "-",
                    "part_code": str(batterypartcode or ""),
                    "qty_system": str(body.quantity), "rate_system": str(unit_price),
                    "price_system": str(quote_price),
                    "rack_dim": "", "qty": "", "per_rack_price": "", "price": "",
                    "custom_cost_desc": "", "custom_cost_price": "",
                    "datasheet": "NO", "sizing_sheet": "YES", "gad": "NO",
                    "battery_compliance": "NO", "warranty": "5 year",
                    "remarks": "", "solution_by": "", "entry_by": "", "data_upload_by": "",
                    "submission_date": "", "submitted_to": "",
                    "quote_code": code, "sol_no": str(sol_no),
                }, db_path=user_inq_db)
        except Exception:
            pass

    return {"detail": "added", "sr_no": sr_no, "quote_price": quote_price}


# ── add from wizard (direct costing data, no tree lookup) ─────────────────────

class AddFromWizardReq(BaseModel):
    battery_config: str = ""
    duration: str = ""
    backup_time_min: str = "0"
    ageing_type: str = "BOL"
    kw_calculation: float = 0
    cell_type: str = ""
    centre_tap: str = ""
    partcode: str = ""
    total_cost: float = 0
    price_option: str = "B"
    quantity: int = 1
    custom_pct: float = 0.0
    actual_load_kva: float = 0
    actual_load_kw: float = 0
    ups_rating_kva: float = 0
    calculated_load_kw: float = 0
    sizing_project: str = ""
    sizing_sr_no: int = 0

@router.post("/quotes/{code}/add-from-sizing-screen", status_code=201)
def add_from_wizard(code: str, body: AddFromWizardReq, user=Depends(get_current_user), scope: str = Query("regular")):
    multipliers = {"A": 1.0, "B-5": 1.05, "B": 1.10, "B+5": 1.15, "C": 1.20, "C+5": 1.25}
    mult = (1.0 + body.custom_pct / 100.0) if body.price_option == "custom" else multipliers.get(body.price_option, 1.10)
    quote_price = round(body.total_cost * mult, 2)

    tdb = _tdb(user["username"], scope)
    quotes = get_all_quotes(tdb)
    meta = next((q for q in quotes if q[0] == code), None)
    if not meta:
        raise HTTPException(404, "Quote not found")
    q_code, q_date, q_customer, q_provider, q_fmt, *_rest = meta; q_sales = _rest[0] if _rest else ""

    sr_no = get_highest_sr_no(code, tdb) + 1
    all_items = get_all_quote_products(code, tdb)
    sol_no = sum(1 for i in all_items if str(_row_to_dict(i)["modular_rack"]) == "-") + 1

    ups_rating_val = str(body.ups_rating_kva) if body.ups_rating_kva > 0 else "-"

    if body.actual_load_kva > 0:
        calc_load_val = str(body.actual_load_kva)
        calc_load_unit = "KVA"
    elif body.actual_load_kw > 0:
        calc_load_val = str(body.actual_load_kw)
        calc_load_unit = "kW"
    else:
        calc_load_val = ""
        calc_load_unit = "kW"

    # backup_req = user-requested time (duration input during sizing)
    raw_req = body.duration
    try:
        backup_req = str(math.floor(float("".join(c for c in raw_req if c.isdigit() or c == ".")))) if raw_req else "-"
    except Exception:
        backup_req = "-"

    # backup_time = calculated actual backup time from sizing output
    raw_bt = body.backup_time_min if body.backup_time_min and body.backup_time_min != "0" else ""
    try:
        backup_time = str(math.floor(float("".join(c for c in raw_bt if c.isdigit() or c == ".")))) if raw_bt else "-"
    except Exception:
        backup_time = "-"

    ageing_type = body.ageing_type or "BOL"

    add_product_quote(
        code, q_code, q_fmt, q_date, q_provider, q_customer,
        sr_no, sol_no, ups_rating_val,
        backup_req, calc_load_val,
        str(body.cell_type), str(body.centre_tap), str(body.partcode),
        backup_time, body.quantity, quote_price, "-",
        calc_load_unit=calc_load_unit, item_type="system", ageing_type=ageing_type, db_path=tdb,
    )

    try:
        import time as _time, sqlite3 as _sq3, re as _re
        from inquiry_db import push_row as _push_inq
        _yr = _time.localtime().tm_year % 100
        _type = f"EVTPL/{_yr:02d}-{(_yr+1):02d}/{code}"
        unit_price = round(quote_price / body.quantity, 2) if body.quantity else quote_price
        user_inq_db = get_user_inquiry_db(user["username"])

        if body.sizing_project and body.sizing_sr_no:
            sdb = get_user_sizing_db(user["username"])
            srow = fetch_sizing_by_sr(body.sizing_project, body.sizing_sr_no, db_path=sdb)
            if srow:
                _srow_ageing = str(srow[31]) if len(srow) > 31 and srow[31] else "BOL"
                _srow_bt = str(math.floor(float(srow[30]))) if len(srow) > 30 and srow[30] else "-"
                _push_inq({
                    "inquiry_date": q_date,
                    "type": _type, "sales_person": str(q_sales or ""),
                    "solution_provider": str(q_provider or ""),
                    "project_customer": str(q_customer or ""),
                    "ups_make": str(srow[3] or ""), "ups_model": str(srow[4] or ""),
                    "ups_kva": str(srow[5] or ""),
                    "actual_load_kva": str(srow[6] or ""), "load_kw": str(srow[7] or ""),
                    "power_factor": str(srow[8] or ""), "inverter_efficiency": str(srow[9] or ""),
                    "dc_voltage": str(srow[10] or ""), "backup_min": str(srow[11] or ""),
                    "cell_chemistry": str(srow[17] or ""),
                    "ageing_pct": str(srow[12] or ""), "design_margin_pct": str(srow[13] or ""),
                    "dod_margin_pct": str(srow[14] or ""), "derating_pct": str(srow[15] or ""),
                    "capacity_ah": str(srow[27] or ""),
                    "centre_tap": str(body.centre_tap or ""), "cell_type": str(body.cell_type or ""),
                    "ageing_type": _srow_ageing, "backup_time_min": _srow_bt,
                    "part_code": str(body.partcode or ""),
                    "qty_system": str(body.quantity), "rate_system": str(unit_price),
                    "price_system": str(quote_price),
                    "rack_dim": "", "qty": "", "per_rack_price": "", "price": "",
                    "custom_cost_desc": "", "custom_cost_price": "",
                    "datasheet": "NO", "sizing_sheet": "YES", "gad": "NO",
                    "battery_compliance": "NO", "warranty": "5 year",
                    "remarks": "", "solution_by": "", "entry_by": "", "data_upload_by": "",
                    "submission_date": "", "submitted_to": "",
                    "quote_code": code, "sol_no": str(sol_no),
                }, db_path=user_inq_db)
        else:
            # wizard flow — use body data + optional costing row lookup by partcode
            _dc_volt = ""
            _volt_m = _re.search(r'(\d+)\s*[Vv]', body.battery_config or "")
            if _volt_m:
                _dc_volt = _volt_m.group(1)
            _crow: dict = {}
            if body.partcode:
                try:
                    _cdb = get_user_costing_db(user["username"])
                    _cconn = _sq3.connect(_cdb)
                    _ccur = _cconn.cursor()
                    _ccur.execute("PRAGMA table_info(tree)")
                    _ccols = [r[1] for r in _ccur.fetchall()]
                    _ccur.execute("SELECT * FROM tree WHERE partcode = ?", (body.partcode,))
                    _crow_row = _ccur.fetchone()
                    if _crow_row:
                        _crow = dict(zip(_ccols, _crow_row))
                    _cconn.close()
                except Exception:
                    pass
            _push_inq({
                "inquiry_date": q_date or _time.strftime("%d/%m/%Y"),
                "type": _type, "sales_person": str(q_sales or ""),
                "solution_provider": str(q_provider or ""),
                "project_customer": str(q_customer or ""),
                "ups_make": "", "ups_model": "",
                "ups_kva": str(body.ups_rating_kva) if body.ups_rating_kva else "",
                "actual_load_kva": str(body.actual_load_kva) if body.actual_load_kva else "",
                "load_kw": str(body.actual_load_kw or body.calculated_load_kw or ""),
                "power_factor": "", "inverter_efficiency": "",
                "dc_voltage": _dc_volt,
                "backup_min": str(backup_req if backup_req != "-" else ""),
                "cell_chemistry": str(_crow.get("cell_chemistry", "") or "LFP"),
                "ageing_pct": "", "design_margin_pct": "", "dod_margin_pct": "", "derating_pct": "",
                "capacity_ah": str(_crow.get("ampere_capacity", "") or ""),
                "centre_tap": str(body.centre_tap or ""), "cell_type": str(body.cell_type or ""),
                "ageing_type": ageing_type, "backup_time_min": backup_time,
                "part_code": str(body.partcode or ""),
                "qty_system": str(body.quantity), "rate_system": str(unit_price),
                "price_system": str(quote_price),
                "rack_dim": "", "qty": "", "per_rack_price": "", "price": "",
                "custom_cost_desc": "", "custom_cost_price": "",
                "datasheet": "NO", "sizing_sheet": "NO", "gad": "NO",
                "battery_compliance": "NO", "warranty": "5 year",
                "remarks": "", "solution_by": "", "entry_by": "", "data_upload_by": "",
                "submission_date": "", "submitted_to": "",
                "quote_code": code, "sol_no": str(sol_no),
            }, db_path=user_inq_db)
    except Exception:
        pass

    return {"detail": "added", "sr_no": sr_no, "quote_price": quote_price}


# ── add modular rack ──────────────────────────────────────────────────────────

@router.post("/quotes/{code}/add-modular", status_code=201)
def add_modular(code: str, body: AddModularReq, user=Depends(get_current_user)):
    if body.custom_price > 0:
        price = body.custom_price
    else:
        price = MODULAR_RACKS.get(body.rack_key)
        if price is None:
            raise HTTPException(400, f"Unknown rack: {body.rack_key}")

    tdb = get_user_temp_db(user["username"])
    quotes = get_all_quotes(tdb)
    meta = next((q for q in quotes if q[0] == code), None)
    if not meta:
        raise HTTPException(404, "Quote not found")
    q_code, q_date, q_customer, q_provider, q_fmt, *_rest = meta; q_sales = _rest[0] if _rest else ""

    sr_no = get_highest_sr_no(code, tdb) + 1
    add_product_quote(
        code, q_code, q_fmt, q_date, q_provider, q_customer,
        sr_no, "-", "-", "-", "-", "-", "-", "-", "-",
        body.quantity, price, body.rack_key, item_type="rack", db_path=tdb,
    )
    try:
        from inquiry_db import sync_inquiry_for_quote as _sync_inq
        user_inq_db = get_user_inquiry_db(user["username"])
        updated = [_row_to_dict(i) for i in get_all_quote_products(code, tdb)]
        _sync_inq(code, updated, db_path=user_inq_db)
    except Exception:
        pass
    return {"detail": "added", "sr_no": sr_no, "quote_price": price}


# ── add custom cost ───────────────────────────────────────────────────────────

class AddCustomCostReq(BaseModel):
    description: str
    price: float
    quantity: int = 1

@router.post("/quotes/{code}/add-custom-cost", status_code=201)
def add_custom_cost(code: str, body: AddCustomCostReq, user=Depends(get_current_user)):
    tdb = get_user_temp_db(user["username"])
    quotes = get_all_quotes(tdb)
    meta = next((q for q in quotes if q[0] == code), None)
    if not meta:
        raise HTTPException(404, "Quote not found")
    q_code, q_date, q_customer, q_provider, q_fmt, *_rest = meta; q_sales = _rest[0] if _rest else ""

    sr_no = get_highest_sr_no(code, tdb) + 1
    add_product_quote(
        code, q_code, q_fmt, q_date, q_provider, q_customer,
        sr_no, "-", "-", "-", "-", "-", "-", "-", "-",
        body.quantity, body.price, body.description, item_type="custom", db_path=tdb,
    )
    try:
        from inquiry_db import sync_inquiry_for_quote as _sync_inq
        user_inq_db = get_user_inquiry_db(user["username"])
        updated = [_row_to_dict(i) for i in get_all_quote_products(code, tdb)]
        _sync_inq(code, updated, db_path=user_inq_db)
    except Exception:
        pass
    return {"detail": "added", "sr_no": sr_no, "quote_price": body.price}


# ── Firebase quote sync ───────────────────────────────────────────────────────

@router.get("/firebase-quotes")
def firebase_list(_=Depends(get_current_user)):
    try:
        from basefunctions import fetch_quote_info
        return fetch_quote_info()
    except Exception as e:
        raise HTTPException(500, str(e))


@router.post("/quotes/{code}/save-to-firebase")
def save_to_firebase(code: str, user=Depends(get_current_user)):
    try:
        from basefunctions import save_quote
        tdb = get_user_temp_db(user["username"])
        save_quote(target_code=code, db_path=tdb)
    except Exception as e:
        raise HTTPException(500, str(e))
    return {"detail": "saved"}


@router.post("/firebase-quotes/{code}/download", status_code=201)
def download_from_firebase(code: str, user=Depends(get_current_user)):
    try:
        from basefunctions import download_quote_from_firebase
        data = download_quote_from_firebase(code)
    except Exception as e:
        raise HTTPException(500, str(e))
    if not data:
        raise HTTPException(404, "Quote not found in Firebase")

    items_list = list(data.values()) if isinstance(data, dict) else data
    if not items_list:
        raise HTTPException(404, "Empty quote")
    first = items_list[0] if isinstance(items_list[0], dict) else {}
    q_code = first.get("code", code)
    q_date = first.get("date", datetime.now().strftime("%d/%m/%Y"))
    q_customer = first.get("customer_name", "")
    q_provider = first.get("solution_provider", "")
    q_fmt = first.get("format", "Quote_format_High_Vtg.docx")

    tdb = get_user_temp_db(user["username"])
    init_temp_db(tdb)
    add_new_quote(q_code, q_date, q_customer, q_provider, q_fmt, tdb)
    clear_quotedata_table(q_code, tdb)

    sr_no = 0
    for item in items_list:
        if not isinstance(item, dict):
            continue
        sr_no += 1
        add_product_quote(
            q_code, item.get("code", q_code), item.get("format", q_fmt),
            item.get("date", q_date), item.get("solution_provider", q_provider),
            item.get("customer_name", q_customer),
            sr_no, item.get("sol_no", "-"),
            item.get("ups_rating", "-"), item.get("backup_requirement", "-"),
            item.get("calc_load", "-"), item.get("celltype", "-"),
            item.get("centre_tapping", "-"), item.get("batterypartcode", "-"),
            item.get("backup_time", "0"), item.get("quantity", 1),
            item.get("quote_price", 0), item.get("modular_rack", "-"), db_path=tdb,
        )
    return {"code": q_code}


# ── Update item ───────────────────────────────────────────────────────────────

class UpdateItemReq(BaseModel):
    ups_rating: Optional[str] = None
    backup_requirement: Optional[str] = None
    calc_load: Optional[str] = None
    celltype: Optional[str] = None
    centre_tapping: Optional[str] = None
    batterypartcode: Optional[str] = None
    backup_time: Optional[str] = None
    quantity: Optional[int] = None
    quote_price: Optional[float] = None
    system_text: Optional[str] = None
    solution_text: Optional[str] = None

@router.patch("/quotes/{code}/items/{sr_no}")
def update_item(code: str, sr_no: int, body: UpdateItemReq, user=Depends(get_current_user), scope: str = Query("regular")):
    from tempquotebase import get_items_table_name, get_db_connection
    tdb = _tdb(user["username"], scope)
    table = get_items_table_name(code)
    fields = {k: v for k, v in body.model_dump().items() if v is not None}
    if not fields:
        return {"detail": "nothing to update"}
    set_clause = ", ".join(f'"{k}" = ?' for k in fields)
    values = list(fields.values()) + [sr_no]
    conn = get_db_connection(tdb)
    try:
        for col in ["system_text", "solution_text"]:
            try:
                conn.execute(f'ALTER TABLE "{table}" ADD COLUMN "{col}" text')
            except Exception:
                pass
        conn.execute(f'UPDATE "{table}" SET {set_clause} WHERE sr_no = ?', values)
        conn.commit()
    except Exception as e:
        conn.rollback()
        raise HTTPException(500, str(e))
    finally:
        conn.close()
    return {"detail": "updated"}


# ── Reorder ───────────────────────────────────────────────────────────────────

class ReorderReq(BaseModel):
    sr_nos: list[int]

@router.put("/quotes/{code}/reorder")
def reorder_items(code: str, body: ReorderReq, user=Depends(get_current_user), scope: str = Query("regular")):
    import sqlite3
    from tempquotebase import get_items_table_name, get_db_connection
    tdb = _tdb(user["username"], scope)
    table = get_items_table_name(code)
    conn = get_db_connection(tdb)
    c = conn.cursor()
    try:
        for i, sr in enumerate(body.sr_nos):
            c.execute(f'UPDATE "{table}" SET sr_no = ? WHERE sr_no = ?', (-(i + 1), sr))
        for i in range(len(body.sr_nos)):
            c.execute(f'UPDATE "{table}" SET sr_no = ? WHERE sr_no = ?', (i + 1, -(i + 1)))
        conn.commit()
    except Exception as e:
        conn.rollback()
        raise HTTPException(500, str(e))
    finally:
        conn.close()
    try:
        from inquiry_db import sync_inquiry_for_quote as _sync_inq
        updated = [_row_to_dict(i) for i in get_all_quote_products(code, tdb)]
        _sync_inq(code, updated)
    except Exception:
        pass
    return {"detail": "reordered"}


# ── Export ────────────────────────────────────────────────────────────────────

@router.get("/quotes/{code}/export/word")
def export_word(code: str, scope: str = Query("regular"), user=Depends(get_current_user)):
    tdb = _tdb(user["username"], scope)
    try:
        path = _generate_docx(code, tdb)
    except Exception as e:
        raise HTTPException(500, str(e))

    try:
        from inquiry_db import push_to_global as _push_to_global
        user_inq_db = get_user_inquiry_db(user["username"])
        _push_to_global(code, user_inq_db)
    except Exception:
        pass

    fname = f"Quote_{code}.docx"
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=fname,
    )


@router.get("/quotes/{code}/export/pdf")
def export_pdf(code: str, scope: str = Query("regular"), user=Depends(get_current_user)):
    tdb = _tdb(user["username"], scope)
    try:
        docx_path = _generate_docx(code, tdb)
        pdf_path = docx_path.replace(".docx", ".pdf")
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close()
        finally:
            word.Quit()
        os.unlink(docx_path)
    except ImportError:
        raise HTTPException(501, "PDF export requires Microsoft Word installed on the server")
    except Exception as e:
        raise HTTPException(500, str(e))
    try:
        from inquiry_db import push_to_global as _push_to_global
        user_inq_db = get_user_inquiry_db(user["username"])
        _push_to_global(code, user_inq_db)
    except Exception:
        pass
    return FileResponse(pdf_path, media_type="application/pdf", filename=f"Quote_{code}.pdf")


# ── record restore ────────────────────────────────────────────────────────────

class RestoreQuoteItem(BaseModel):
    sr_no: int = 1
    sol_no: int = 1
    ups_rating: str = "-"
    backup_requirement: str = "-"
    calc_load: str = "-"
    celltype: str = "-"
    centre_tapping: str = "-"
    batterypartcode: str = "-"
    backup_time: str = "0"
    quantity: int = 1
    quote_price: float = 0
    modular_rack: str = "-"

class RestoreQuoteMeta(BaseModel):
    code: str = ""
    date: str = ""
    customer_name: str = ""
    solution_provider: str = ""
    format_name: str = "High voltage"

class RestoreQuoteReq(BaseModel):
    meta: RestoreQuoteMeta
    items: list[RestoreQuoteItem]


@router.post("/restore")
def restore_quote(body: RestoreQuoteReq, user=Depends(get_current_user)):
    meta = body.meta
    fname = TEMPLATE_MAP.get(meta.format_name, meta.format_name) or "Quote_format_High_Vtg.docx"

    tdb = get_user_temp_db(user["username"])
    init_temp_db(tdb)
    existing_codes = {q[0] for q in get_all_quotes(tdb)}
    code = meta.code
    if code in existing_codes:
        from datetime import datetime as _dt
        suffix = _dt.now().strftime("%m%d%H%M")
        code = f"{meta.code}-R{suffix}"

    try:
        add_new_quote(code, meta.date, meta.customer_name, meta.solution_provider, fname, tdb)
        for item in body.items:
            add_product_quote(
                code, code, fname, meta.date,
                meta.solution_provider, meta.customer_name,
                item.sr_no, item.sol_no, item.ups_rating,
                item.backup_requirement, item.calc_load, item.celltype,
                item.centre_tapping, item.batterypartcode, item.backup_time,
                item.quantity, item.quote_price, item.modular_rack, db_path=tdb,
            )
    except Exception as e:
        raise HTTPException(500, str(e))
    return {"code": code}
